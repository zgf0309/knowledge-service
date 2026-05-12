# -*- coding: utf-8 -*-
"""
file-service 路由
对外接口与 jusure_AI 对齐：
  POST /ai/files/upload        — 文件上传
  GET  /ai/knowledge/mdcontent — 获取文档 MD 内容（按 object_key）
  GET  /ai/files/url           — 获取预签名 URL
  DELETE /ai/files             — 删除文件
  POST /ai/knowledge/file2md   — 文件转 MD（批量，传 oss_names_array）[高优先级]
"""
import os
from typing import Any

from fastapi import APIRouter, UploadFile, File, Query, Body, Depends, HTTPException
from fastapi.responses import Response

from common.auth_context import RequestContext, get_request_context, pick_tenant
from common.utils import get_logger
from common.utils.response import api_success

from .services import FileService

logger = get_logger("file_api")
router = APIRouter(prefix="/ai", tags=["File Service"])

_file_service = FileService()

# ---------------------------------------------------------------------------
# POST /ai/files/upload
# 对应 jusure_AI: MinIOUploadView POST /ai/files/upload
# 参考 ragflow: api/apps/file_app.py upload
# ---------------------------------------------------------------------------
@router.post("/files/upload")
async def upload_file(
    file: UploadFile = File(...), kb_id: str | None = Query(None, description="知识库ID（可选）"), ctx: RequestContext = Depends(get_request_context), ):
    """
    上传文件到对象存储（MinIO）
    返回 doc_url（预签名 URL）和 object_key（内部路径）
    """
    try:
        content = await file.read()
        object_key, doc_url = await _file_service.upload_file(
            tenant_id=pick_tenant(ctx), file_data=content, filename=file.filename or "unknown", kb_id=kb_id, content_type=file.content_type, )
        return api_success(data={
            "doc_url": doc_url, "object_key": object_key, "doc_name": file.filename, "doc_size": len(content), })
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.exception("upload_file error")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# GET /ai/knowledge/mdcontent
# 对应 jusure_AI: GetMdContentView GET /ai/knowledge/mdcontent
# 获取已上传文档的 MD 内容
# ---------------------------------------------------------------------------
@router.get("/knowledge/mdcontent")
async def get_md_content(
    object_key: str = Query(..., description="对象存储路径或完整MinIO URL"), ):
    """获取文档内容（支持简单路径或完整MinIO URL）
    
    对于 .md/.txt 文件直接返回内容；
    对于 .docx 文件自动转换为 Markdown；
    其他格式暂不支持转换，返回错误提示。
    """
    try:
        # 解析 object_key，支持两种格式：
        # 1. 简单路径：default/tmp/xxx.md
        # 2. 完整 MinIO URL：http://localhost:9000/jusure/default/tmp/xxx.md?X-Amz-...
        actual_key = object_key
        
        if object_key.startswith("http://") or object_key.startswith("https://"):
            # 解析完整 URL，提取 object_key
            from urllib.parse import urlparse, unquote
            try:
                parsed = urlparse(object_key)
                # 路径格式：/jusure/default/tmp/xxx.md
                path = parsed.path.lstrip('/')  # 去掉开头的 /
                # 去掉 bucket 前缀
                parts = path.split('/', 1)
                if len(parts) >= 2:
                    actual_key = parts[1]  # default/tmp/xxx.md
                else:
                    actual_key = path
                # URL 解码（处理 %28 -> ( 等编码字符）
                actual_key = unquote(actual_key)
            except Exception as e:
                logger.warning(f"URL 解析失败，使用原值：{e}")
        
        logger.info(f"获取文档内容，object_key={actual_key}")
        data = await _file_service.get_file_content(actual_key)
        
        # 判断文件类型并处理
        file_name = actual_key.split('/')[-1].lower()
        
        if file_name.endswith('.md') or file_name.endswith('.txt'):
            # 文本文件直接解码
            text = data.decode("utf-8", errors="replace")
            return api_success(data={"content": text, "object_key": actual_key})
        
        elif file_name.endswith('.docx'):
            # DOCX 文件转换为 Markdown
            try:
                from docx import Document
                from io import BytesIO
                
                doc = Document(BytesIO(data))
                md_lines: list[str] = []
                
                # 段落
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = para.style.name if para.style else ""
                    if style_name.startswith("Heading"):
                        try:
                            level = int(style_name.replace("Heading", "").strip())
                        except ValueError:
                            level = 1
                        md_lines.append(f"{'#' * level} {text}")
                    else:
                        md_lines.append(text)
                
                # 表格
                for table in doc.tables:
                    md_lines.append("")
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        md_lines.append("| " + " | ".join(cells) + " |")
                    md_lines.append("")
                
                markdown_content = "\n".join(md_lines)
                return api_success(data={"content": markdown_content, "object_key": actual_key})
            except ImportError:
                raise HTTPException(status_code=500, detail="DOCX 解析依赖未安装：python-docx")
            except Exception as e:
                logger.exception("DOCX 转 Markdown 失败")
                raise HTTPException(status_code=500, detail=f"DOCX 解析失败：{str(e)}")
        
        elif file_name.endswith(('.mp3', '.wav', '.m4a', '.flac', '.ogg', '.wma')):
            # 音频文件：返回元数据，转写文本需通过 media_service 获取
            return api_success(data={
                "content": "", "object_key": actual_key, "file_type": "audio", "audio_format": file_name.rsplit('.', 1)[-1], "message": "音频文件内容需通过语音转录接口获取，请调用 /ai/media/transcription 接口"
            })

        elif file_name.endswith(('.mp4', '.avi', '.mov', '.flv', '.wmv', '.mkv')):
            # 视频文件：返回元数据
            return api_success(data={
                "content": "", "object_key": actual_key, "file_type": "video", "message": "视频文件内容需通过媒体服务获取"
            })

        elif file_name.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            # 图片文件：返回元数据
            return api_success(data={
                "content": "", "object_key": actual_key, "file_type": "image", "message": "图片文件内容需通过相应接口获取"
            })

        else:
            # 其他格式暂不支持
            raise HTTPException(
                status_code=400, detail=f"不支持的文件格式：{file_name}。目前仅支持 .md、.txt、.docx 文件"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("get_md_content error")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# GET /ai/files/url
# 获取预签名下载 URL（expires 默认 3600 秒）
# ---------------------------------------------------------------------------
@router.get("/files/url")
async def get_file_url(
    object_key: str = Query(...), expires: int = Query(3600, ge=60, le=604800), ):
    """获取预签名下载 URL"""
    try:
        url = await _file_service.get_presigned_url(object_key, expires=expires)
        return api_success(data={"url": url, "expires": expires})
    except Exception as e:
        logger.exception("get_file_url error")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# DELETE /ai/files
# 删除对象存储中的文件
# ---------------------------------------------------------------------------
@router.delete("/files")
async def delete_file(
    object_key: str = Query(...), ctx: RequestContext = Depends(get_request_context), ):
    """删除文件"""
    try:
        tenant_id = pick_tenant(ctx)
        # 基本路径校验：object_key 必须以 tenant_id 开头，防止越权
        if not object_key.startswith(f"{tenant_id}/"):
            raise HTTPException(status_code=403, detail="无权删除该文件")
        await _file_service.delete_file(object_key)
        return api_success(message="删除成功")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("delete_file error")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# POST /ai/knowledge/file2md  [高优先级]
# 对应 jusure_AI File2MdView POST /ai/knowledge/file2md
# 将对象存储中的文件转换为 Markdown 内容（批量）
# ---------------------------------------------------------------------------
@router.post("/knowledge/file2md")
async def file_to_md(
    body: dict[str, Any] = Body(...), ctx: RequestContext = Depends(get_request_context), ):
    """将对象存储中的文件批量转换为 Markdown 内容

    对应 jusure_AI File2MdView POST /ai/knowledge/file2md
    body: { oss_names_array: ["object/key1", "object/key2", ...] }
    """
    oss_names = body.get("oss_names_array", [])
    if not oss_names:
        raise HTTPException(status_code=400, detail="oss_names_array 不能为空")

    tenant_id = pick_tenant(ctx)
    results: list[dict[str, Any]] = []
    audio_exts = {'.mp3', '.wav', '.m4a', '.flac', '.ogg', '.wma'}
    video_exts = {'.mp4', '.avi', '.mov', '.flv', '.wmv', '.mkv'}
    image_exts = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}
    for object_key in oss_names:
        try:
            raw = await _file_service.get_file_content(object_key)
            ext = os.path.splitext(object_key)[1].lower()

            if ext in audio_exts:
                results.append({
                    "object_key": object_key, "content": "", "success": True, "file_type": "audio", "message": "音频文件需通过语音转录接口获取文本内容", })
            elif ext in video_exts:
                results.append({
                    "object_key": object_key, "content": "", "success": True, "file_type": "video", "message": "视频文件需通过媒体服务获取内容", })
            elif ext in image_exts:
                results.append({
                    "object_key": object_key, "content": "", "success": True, "file_type": "image", "message": "图片文件需通过相应接口获取内容", })
            else:
                # 文本类文件，尝试 UTF-8 解码
                try:
                    text = raw.decode("utf-8")
                except UnicodeDecodeError:
                    text = raw.decode("utf-8", errors="replace")
                results.append({
                    "object_key": object_key, "content": text, "success": True, })
        except Exception as e:
            logger.warning(f"file2md failed for {object_key}: {e}")
            results.append({
                "object_key": object_key, "content": "", "success": False, "error": str(e), })

    return api_success(data={"results": results, "total": len(results)})
